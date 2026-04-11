from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
ACCENT   = "03363D"
ACCENT2  = "78D64B"
WHITE    = "FFFFFF"
LIGHT_BG = "F4F6F8"
MID_GRAY = "4A5568"
BORDER_C = "CBD5E0"

SS = "/home/fyxvoid/void/ceh/zendesk/screenshot/"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc       = cell._tc
    tcPr     = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = kwargs.get(edge, {})
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   val.get("val",   "single"))
        tag.set(qn("w:sz"),    val.get("sz",    "4"))
        tag.set(qn("w:space"), val.get("space", "0"))
        tag.set(qn("w:color"), val.get("color", BORDER_C))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def row_border(row):
    for cell in row.cells:
        b = {"val": "single", "sz": "4", "color": BORDER_C}
        set_cell_border(cell, top=b, bottom=b, left=b, right=b)

def para_space(para, before=0, after=0):
    pPr  = para._p.get_or_add_pPr()
    spng = OxmlElement("w:spacing")
    spng.set(qn("w:before"), str(before))
    spng.set(qn("w:after"),  str(after))
    pPr.append(spng)

def add_run(para, text, bold=False, italic=False,
            size=11, color=MID_GRAY, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def heading(doc, text, level=1):
    sizes  = {1: 16, 2: 13, 3: 11}
    colors = {1: ACCENT, 2: ACCENT, 3: MID_GRAY}
    p = doc.add_paragraph()
    para_space(p, before=160, after=60)
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    "24")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), ACCENT2)
        pBdr.append(left)
        pPr.append(pBdr)
    add_run(p, text, bold=True, size=sizes[level],
            color=colors[level], font="Calibri Light")
    return p

def build_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, ACCENT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_space(p, before=40, after=40)
        add_run(p, h, bold=True, size=9.5, color=WHITE)

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg  = LIGHT_BG if r_idx % 2 == 0 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            para_space(p, before=30, after=30)
            add_run(p, str(cell_text), size=9.5, color=MID_GRAY)
            row_border(row)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return tbl

def code_block(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=40, after=40)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EDF2F7")
    pPr.append(shd)
    add_run(p, text, size=9, color="2D3748", font="Courier New")
    return p

def body(doc, text, size=10.5, color=MID_GRAY, italic=False):
    p = doc.add_paragraph()
    para_space(p, before=20, after=20)
    add_run(p, text, size=size, color=color, italic=italic)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=40, after=40)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ACCENT2)
    pBdr.append(left)
    pPr.append(pBdr)
    add_run(p, "Note: ", bold=True, size=10, color=ACCENT)
    add_run(p, text, size=10, color=MID_GRAY, italic=True)
    return p

def add_image(doc, filename, caption=None, width=Inches(5.8)):
    path = os.path.join(SS, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p, before=40, after=20)
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_space(cap, before=0, after=60)
        add_run(cap, f"Figure: {caption}", italic=True, size=9, color="718096")

# ─────────────────────────────────────────────────────────────────────────────
#  COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(cover, before=400, after=20)
add_run(cover, "OSINT RECONNAISSANCE REPORT",
        bold=True, size=26, color=ACCENT, font="Calibri Light")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(sub, before=0, after=80)
add_run(sub, "Target: zendesk.com",
        bold=False, size=15, color=MID_GRAY, font="Calibri Light")

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate([
    ("Author",  "sirdharan thangaraji"),
    ("Date",    "2026-03-31"),
    ("Scope",   "Passive OSINT — zendesk.com"),
    ("Method",  "Passive / Open-Source Intelligence"),
]):
    bg = LIGHT_BG if i % 2 == 0 else WHITE
    kc = meta.rows[i].cells[0]
    vc = meta.rows[i].cells[1]
    set_cell_bg(kc, bg); set_cell_bg(vc, bg)
    row_border(meta.rows[i])
    para_space(kc.paragraphs[0], before=40, after=40)
    para_space(vc.paragraphs[0], before=40, after=40)
    kc.width = Inches(2.0); vc.width = Inches(4.0)
    add_run(kc.paragraphs[0], k, bold=True, size=10, color=ACCENT)
    add_run(vc.paragraphs[0], v, size=10,   color=MID_GRAY)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 1 — Executive Summary
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "1. Executive Summary")
body(doc,
     "This report documents passive OSINT reconnaissance conducted against zendesk.com. "
     "No active exploitation was performed. The objective was to map publicly available "
     "information including domain registration data, DNS infrastructure, exposed subdomains, "
     "IP geolocation, network path, web surface metadata, and employee/corporate intelligence.")
body(doc, "Key findings:")
for b in [
    "Domain is protected by MarkMonitor with all EPP lock statuses enabled.",
    "DNS is hosted on AWS Route 53; web traffic routes through Cloudflare (WAF/CDN).",
    "Over 200 subdomains were enumerated via passive sources.",
    "robots.txt reveals internal path structure including marketplace, brand, and partner directories.",
    "Email infrastructure runs on Google Workspace.",
    "LinkedIn enumeration identified Zendesk India employees in Bengaluru.",
    "Indian subsidiary (Zendesk Technologies Pvt. Ltd.) director details are publicly available via MCA records.",
]:
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=10, after=10)
    add_run(p, b, size=10.5, color=MID_GRAY)

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 2 — WHOIS
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "2. Domain Registration (WHOIS)")
build_table(doc,
    ["Field", "Value"],
    [
        ("Domain",               "zendesk.com"),
        ("Registry Domain ID",   "157863981_DOMAIN_COM-VRSN"),
        ("Registrar",            "MarkMonitor Inc. (IANA ID: 292)"),
        ("Created",              "2005-05-16"),
        ("Last Updated",         "2024-04-15"),
        ("Expiry",               "2027-05-16"),
        ("Registrant Org",       "Zendesk, Inc."),
        ("Registrant Country",   "US"),
        ("Registrant Email",     "Redacted (MarkMonitor WHOIS form)"),
        ("DNSSEC",               "Unsigned"),
        ("Abuse Contact",        "abusecomplaints@markmonitor.com"),
    ],
    col_widths=[2.2, 4.0]
)
heading(doc, "Domain Lock Status (all enabled)", level=3)
for lock in [
    "clientDeleteProhibited", "clientTransferProhibited", "clientUpdateProhibited",
    "serverDeleteProhibited", "serverTransferProhibited", "serverUpdateProhibited",
]:
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=8, after=8)
    add_run(p, lock, size=10, color=MID_GRAY)
note(doc, "DNSSEC is not enabled. This leaves the domain theoretically susceptible to DNS cache poisoning if not mitigated elsewhere.")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 3 — DNS Enumeration
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3. DNS Enumeration")

heading(doc, "3.1  A Records (nslookup)", level=2)
code_block(doc, "zendesk.com  →  216.198.53.2\nzendesk.com  →  216.198.54.2")
add_image(doc, "2026-03-31_06-47.png",
          caption="nslookup zendesk.com — resolves to 216.198.53.2 and 216.198.54.2")

heading(doc, "3.2  Name Servers (AWS Route 53)", level=2)
build_table(doc,
    ["Name Server"],
    [
        ("ns-1213.awsdns-23.org",),
        ("ns-1858.awsdns-40.co.uk",),
        ("ns-444.awsdns-55.com",),
        ("ns-754.awsdns-30.net",),
    ],
    col_widths=[4.0]
)
body(doc, "DNS is fully managed via Amazon Route 53, indicating Zendesk's infrastructure is AWS-centric at the DNS layer.")

heading(doc, "3.3  MX Records (Email Infrastructure)", level=2)
build_table(doc,
    ["Priority", "Mail Server"],
    [
        ("Mx0", "aspmx.l.google.com"),
        ("Mx1", "alt1.aspmx.l.google.com"),
        ("Mx2", "alt2.aspmx.l.google.com"),
        ("Mx3", "alt3.aspmx.l.google.com"),
        ("Mx4", "alt4.aspmx.l.google.com"),
    ],
    col_widths=[1.2, 4.0]
)
body(doc, "Zendesk uses Google Workspace for corporate email. SPF and DMARC records are present.")
add_image(doc, "2026-03-31_06-56.png",
          caption="MX / DNS record check — Google Workspace mail servers confirmed")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 4 — IP Intelligence
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "4. IP Intelligence")

heading(doc, "4.1  IP Geolocation (216.198.54.2)", level=2)
build_table(doc,
    ["Field", "Value"],
    [
        ("IP Address",    "216.198.54.2"),
        ("Country",       "United States"),
        ("Region",        "California"),
        ("City",          "San Francisco"),
        ("Organization",  "AS209242 — Cloudflare, LLC"),
    ],
    col_widths=[2.2, 3.8]
)
add_image(doc, "2026-03-31_06-48.png",
          caption="ipinfo.io — Geolocation data for 216.198.54.2 (Cloudflare / San Francisco)")
body(doc, "The resolved IPs are Cloudflare edge nodes, meaning the true origin server IP is hidden behind Cloudflare's reverse proxy.")

heading(doc, "4.2  Port Scan / Technology Stack (Shodan)", level=2)
body(doc, 'Cloudflare enforces "Direct IP access not allowed", blocking direct origin access. '
         "Port scan data reflects Cloudflare's edge, not Zendesk's backend. "
         "Web technology fingerprinting confirms CDN: Cloudflare.")
add_image(doc, "2026-03-31_06-49_1.png",
          caption="Shodan — 216.198.54.2 open ports and Cloudflare WAF block message")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 5 — Traceroute
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "5. Network Path (Traceroute — TCP 443)")
build_table(doc,
    ["Hop", "RTT", "Address", "Note"],
    [
        ("1", "3.48 ms",  "192.168.1.1",                       "Local gateway"),
        ("2", "5.60 ms",  "blr-tdc-bngs-02 (103.57.86.19)",    "ISP — Bengaluru"),
        ("3", "8.24 ms",  "103.57.86.1",                       "ISP uplink"),
        ("4", "—",        "* (no response)",                    "Filtered hop"),
        ("5", "9.02 ms",  "216.198.54.2",                      "Cloudflare edge"),
    ],
    col_widths=[0.5, 1.0, 2.8, 2.0]
)
add_image(doc, "2026-03-31_06-53_1.png",
          caption="Traceroute (TCP 443) — 5-hop path from Bengaluru to Cloudflare edge at 216.198.54.2")
body(doc, "Traffic terminates at Cloudflare's edge in under 10ms from Bengaluru. Origin hops beyond Cloudflare are hidden.")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 6 — Web Surface
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "6. Web Surface Intelligence")

heading(doc, "6.1  robots.txt Analysis", level=2)
body(doc, "The robots.txt at https://www.zendesk.com/robots.txt reveals internal path structure:")
add_image(doc, "2026-03-31_06-50.png",
          caption="robots.txt — disallowed paths exposing internal directory structure")
build_table(doc,
    ["Disallowed Path", "Significance"],
    [
        ("/company/index/thank-you/",             "Post-conversion landing page"),
        ("/zendesk-adwords-template/",            "Ad campaign template path"),
        ("/brand/, /brand",                       "Brand assets directory"),
        ("/public/assets/html/",                  "Public HTML assets"),
        ("/company/add-ons-agreement/",           "Legal/contract pages"),
        ("/company/reseller-add-ons-agreement/",  "Reseller contract pages"),
        ("/public/assets/sitemaps/",              "Additional sitemap location"),
        ("/marketplace/apps",                     "Marketplace app listings"),
        ("/marketplace/partners",                 "Partner directory"),
        ("/marketplace/themes",                   "Theme marketplace"),
        ("/blog/ads-api-output/",                 "Ads API output"),
        ("/blog/tags/*-zip (multiple)",           "Downloadable content archives"),
    ],
    col_widths=[3.0, 3.3]
)

heading(doc, "6.2  Sitemap", level=2)
code_block(doc,
    "Sitemap: https://www.zendesk.com/generated_sitemap.xml\n"
    "  ├── generated_sitemap_part1.xml\n"
    "  └── generated_sitemap_part2.xml"
)
add_image(doc, "2026-03-31_06-50_1.png",
          caption="Sitemap index XML — two-part sitemap structure")
note(doc, "robots.txt effectively acts as a directory disclosure artifact, revealing paths that "
         "Zendesk does not want indexed but are visible to anyone reading the file.")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 7 — Subdomains
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "7. Subdomain Enumeration")
body(doc, "200+ subdomains were passively enumerated. Notable entries by category:")
build_table(doc,
    ["Category", "Subdomains"],
    [
        ("Core",               "www, support, status, developer, training, event, adm"),
        ("Partner/Sales",      "partnersuccess, sellersfundinghelp"),
        ("Security-adjacent",  "offensive-security, contrastsecurity, dune-security, gremlin-security, saltsecurity, siteprotect"),
        ("Crypto/Fintech",     "bitmart, cryptopia, phantom-wallet, safepalsupport, emurgohelpdesk"),
        ("Government/Edu",     "nysed-op, francecompetences, eccouncil"),
        ("Testing/Dev",        "testcenter, nm7uqebiup84rmec6eeb (randomised subdomain)"),
        ("Major Clients",      "gitlab, nvidia1651665403, zillow, wpengine, fourseasons, tescosupportcentre"),
    ],
    col_widths=[2.0, 4.3]
)
note(doc, "The randomised subdomain nm7uqebiup84rmec6eeb.zendesk.com is likely a provisioned test "
         "or dev tenant. Full list in subdomain.txt (200 entries).")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 8 — LinkedIn OSINT
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "8. Employee Intelligence (LinkedIn OSINT)")
add_image(doc, "2026-03-31_06-49.png",
          caption="LinkedIn — Zendesk India employees identified via people search")
build_table(doc,
    ["Name", "Role", "Location"],
    [
        ("Koushik Jain",       "Solutions — Zendesk India",        "Bengaluru"),
        ("Kavitha Dinesh",     "Senior Workplace (1K+ followers)",  "Bengaluru"),
        ("Niharika S N",       "Services Consultant",              "Bengaluru"),
        ("Vishal Srivastava",  "Senior Business Systems",          "Greater Bengaluru"),
    ],
    col_widths=[2.0, 2.8, 1.8]
)
note(doc, "These identities could be targeted in spear-phishing or social engineering scenarios. "
         "Data gathered from public LinkedIn profiles without authentication.")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 9 — Corporate Intelligence
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "9. Corporate Intelligence — Indian Subsidiary")
body(doc, "Entity: Zendesk Technologies Private Limited")
add_image(doc, "2026-03-31_06-50_2.png",
          caption="MCA — Zendesk Technologies Private Limited company overview")
build_table(doc,
    ["Field", "Value"],
    [
        ("Incorporation Date", "19 May 2016"),
        ("Entity Type",        "Foreign Company (India subsidiary)"),
        ("Industry",           "Computer Related Services; Marketing Management Consulting"),
        ("Auth Capital",       "₹20.0 Lakh"),
        ("Paid-up Capital",    "₹5.0 Lakh"),
        ("Revenue",            "₹50–75 Crore"),
        ("Registered Address", "Salarpuria Sigma, Bengaluru, Karnataka"),
    ],
    col_widths=[2.2, 4.0]
)
add_image(doc, "2026-03-31_06-51.png",
          caption="MCA — Company size and financial details")

heading(doc, "Directors", level=3)
add_image(doc, "2026-03-31_06-51_1.png",
          caption="MCA — Director listing with DIN numbers and tenure")
build_table(doc,
    ["Name", "DIN", "Tenure"],
    [
        ("Rajendra Singh Rathore", "02365241", "10 years"),
        ("Julie Ann Swinney",      "10334694", "3 years"),
    ],
    col_widths=[2.8, 1.5, 1.5]
)
note(doc, "Director DIN numbers are public MCA records and can be used to cross-reference other "
         "company affiliations via India's Ministry of Corporate Affairs portal.")

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 10 — Summary of Findings
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "10. Summary of Findings")
build_table(doc,
    ["Category", "Finding", "Risk Level"],
    [
        ("DNSSEC not enabled",      "Domain vulnerable to DNS spoofing if no edge-level mitigation",           "Low–Medium"),
        ("Cloudflare WAF/CDN",      "Origin IP is hidden; direct access blocked",                              "Mitigated"),
        ("robots.txt disclosure",   "Internal path structure exposed",                                         "Informational"),
        ("200+ subdomains",         "Large attack surface; some tenants may be misconfigured",                 "Medium"),
        ("Google Workspace email",  "Target for BEC/phishing if SPF/DMARC not strictly enforced",             "Medium"),
        ("Employee enumeration",    "LinkedIn profiles accessible; social engineering risk",                   "Medium"),
        ("Director PII (India)",    "Publicly available DIN and tenure via MCA",                              "Informational"),
        ("AWS Route 53 DNS",        "DNS infrastructure identified; no immediate risk",                        "Informational"),
    ],
    col_widths=[2.0, 3.5, 1.2]
)

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 11 — Tools Used
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "11. Tools Used")
build_table(doc,
    ["Tool", "Purpose"],
    [
        ("whois",       "Domain registration data"),
        ("nslookup",    "DNS resolution"),
        ("traceroute",  "Network path mapping (TCP 443)"),
        ("ping",        "Subdomain liveness check"),
        ("ipinfo.io",   "IP geolocation"),
        ("Shodan",      "Port scan / web technology detection"),
        ("robots.txt",  "Web surface crawl metadata"),
        ("LinkedIn",    "Employee / personnel OSINT"),
        ("MCA Portal",  "Indian corporate registry lookup"),
        ("Subfinder",   "Passive subdomain enumeration"),
    ],
    col_widths=[1.8, 4.5]
)

# ─────────────────────────────────────────────────────────────────────────────
#  SECTION 12 — Recommendations
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "12. Recommendations")
for i, (title, detail) in enumerate([
    ("Enable DNSSEC",        "Enable DNSSEC on zendesk.com to prevent DNS cache poisoning attacks."),
    ("Audit Subdomains",     "Decommission unused or randomised tenants to reduce the passive attack surface."),
    ("Review robots.txt",    "Evaluate whether disallowed paths expose sensitive structural information that could aid an attacker."),
    ("Enforce DMARC Reject", "Set DMARC policy to p=reject to prevent email spoofing from the zendesk.com domain."),
    ("Employee Awareness",   "LinkedIn-visible personnel are prime social engineering targets; ensure phishing awareness training is in place."),
], 1):
    p = doc.add_paragraph()
    para_space(p, before=20, after=20)
    add_run(p, f"{i}.  {title}: ", bold=True, size=10.5, color=ACCENT)
    add_run(p, detail, size=10.5, color=MID_GRAY)

# ─────────────────────────────────────────────────────────────────────────────
#  Footer disclaimer
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
disc = doc.add_paragraph()
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(disc, before=200, after=0)
add_run(disc,
        "This report is for educational / authorized security assessment purposes only.",
        italic=True, size=9, color="A0AEC0")

# ─────────────────────────────────────────────────────────────────────────────
#  Save
# ─────────────────────────────────────────────────────────────────────────────
out = "/home/fyxvoid/void/ceh/zendesk/osint-report.docx"
doc.save(out)
print(f"Saved → {out}")
