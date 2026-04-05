from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_heading(para, text, level=1, color=RGBColor(0xC0, 0x39, 0x2B)):
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(15)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_heading(p, text, level)
    if level == 1:
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), 'C0392B')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_image(doc, path, caption, width=Inches(5.5)):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        add_caption(doc, caption)
    except Exception as e:
        add_body(doc, f'[Screenshot: {caption}]')

def shade_row(row, hex_color='F2F2F2'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    shade_row(hrow, '1A1A2E')
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = w
    doc.add_paragraph()
    return t

# ════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('CYBERSECURITY FOUNDATION LAB')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run('Technical Assessment Report')
run.bold = True
run.font.size = Pt(28)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Common Assessment — Practical Submission')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)

# Divider line
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '8')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'C0392B')
pBdr.append(bottom)
pPr.append(pBdr)
p.paragraph_format.space_after = Pt(30)

meta = [
    ('Candidate Name', 'Sridharan T'),
    ('Submitted To',   'Mr. Ayush Singh'),
    ('Date',           'April 2026'),
    ('Document',       'Sridharan_T_CommonAssessment.pdf'),
    ('Classification', 'Internal / Lab Use Only'),
    ('Version',        '1.0 — Final'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}:  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════
add_heading(doc, '1. Introduction')
add_body(doc,
    'This report documents the hands-on work completed for the Cybersecurity Foundation Lab Assessment, '
    'assigned by Mr. Ayush Singh. The assessment covers three core practical tasks: network scanning and '
    'enumeration, network traffic analysis, and basic exploitation — all carried out inside an isolated '
    'virtual lab environment.'
)
add_body(doc,
    'The tools used are Nmap for reconnaissance, Wireshark for packet analysis, and the Metasploit '
    'Framework for exploitation. Each task is documented with the actual commands run, screenshots of '
    'the output, and a summary of what was found.'
)

# ════════════════════════════════════════
# 2. LAB SETUP
# ════════════════════════════════════════
add_heading(doc, '2. Lab Setup & Network Configuration')
add_body(doc,
    'The lab was built using VMware Workstation Pro 25H2 on a Windows host. A Host-Only network adapter '
    'was used so all machines are on a private segment with no internet exposure. Kali Linux runs as the '
    'attacker machine (host), Windows 10 x64 as an additional target, and Metasploitable 2 as the '
    'primary intentionally vulnerable target.'
)

add_heading(doc, '2.1  Virtualization Software', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/installation/20260401_084702_flameshot.png',
    'Figure 1 — VMware Workstation Pro 25H2 installed and running',
    width=Inches(5.0)
)

add_heading(doc, '2.2  Virtual Machines', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/installation/20260401_100954_flameshot.png',
    'Figure 2 — All three VMs visible in VMware: Windows 10, Kali Linux, Metasploitable 2',
    width=Inches(5.5)
)

add_heading(doc, '2.3  VM Configuration Details', level=2)
add_table(doc,
    ['Machine', 'Role', 'RAM', 'Disk', 'Network Adapter'],
    [
        ['Windows 10 x64',          'Additional Target',     '2 GB',   '60 GB', 'Host-only'],
        ['Kali Linux 2025.4 amd64', 'Attacker / Host',      '4 GB',   '80 GB', 'Host-only'],
        ['Metasploitable 2',        'Primary Vulnerable Target', '512 MB', '8 GB',  'Host-only'],
    ],
    col_widths=[Inches(1.6), Inches(1.8), Inches(0.7), Inches(0.7), Inches(1.4)]
)

add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/network/20260401_101003_flameshot.png',
    'Figure 3 — Windows 10 VM: Network Adapter set to Host-only',
    width=Inches(3.0)
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/network/20260401_101010_flameshot.png',
    'Figure 4 — Kali Linux VM: Network Adapter set to Host-only',
    width=Inches(3.0)
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/network/20260401_101032_flameshot.png',
    'Figure 5 — Metasploitable 2 VM: Network Adapter set to Host-only',
    width=Inches(3.0)
)

add_heading(doc, '2.4  IP Address Table', level=2)
add_table(doc,
    ['Machine', 'IP Address', 'Interface', 'Role'],
    [
        ['Kali Linux',       '192.168.1.50',   'wlan0',  'Attacker'],
        ['Metasploitable 2', '192.168.38.130', 'eth0',   'Target'],
        ['Windows 10',       '192.168.x.x',    'Ethernet', 'Additional Target'],
    ],
    col_widths=[Inches(1.5), Inches(1.4), Inches(1.1), Inches(1.5)]
)

add_heading(doc, '2.5  Connectivity Verification', level=2)
add_body(doc,
    'Before starting any tasks, connectivity between the machines was confirmed using ping from the '
    'Kali terminal to the Metasploitable target.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/connectivity/20260401_094949_flameshot.png',
    'Figure 6 — Ping from Kali (192.168.1.50) to Metasploitable (192.168.38.130) — 0% packet loss',
    width=Inches(5.5)
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/connectivity/20260401_100546_flameshot.png',
    'Figure 7 — Metasploitable network interface showing IP 192.168.38.130',
    width=Inches(5.5)
)

doc.add_page_break()

# ════════════════════════════════════════
# 3. TASK 1 — NMAP
# ════════════════════════════════════════
add_heading(doc, '3. Task 1 — Network Scanning & Enumeration')
add_body(doc,
    'Nmap was used to discover open ports, identify running services, detect service versions, and '
    'fingerprint the operating system on the Metasploitable target (192.168.38.130). The scan was '
    'done in stages — starting with a ping sweep to confirm the host is live, followed by progressively '
    'deeper scans on specific ports and services.'
)

add_heading(doc, '3.1  Host Discovery', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_102956_flameshot.png',
    'Figure 8 — Nmap ping sweep confirming Metasploitable host is up',
    width=Inches(5.5)
)

add_heading(doc, '3.2  Full Port Scan — All Services', level=2)
add_body(doc,
    'A full TCP scan was run against the target. The output below shows all open ports with their '
    'associated service names.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_104237_flameshot.png',
    'Figure 9 — Nmap full port scan output showing all open ports on Metasploitable',
    width=Inches(3.5)
)

add_heading(doc, '3.3  FTP Service Scan (Port 21)', level=2)
add_body(doc,
    'A detailed scan of port 21 revealed vsftpd 2.3.4 running with anonymous login allowed. '
    'The -A flag was used to get full service and OS detail.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_110250_flameshot.png',
    'Figure 10 — Nmap -sV -sC -A on port 21: vsftpd 2.3.4, anonymous FTP login allowed',
    width=Inches(5.5)
)

add_heading(doc, '3.4  Vulnerability Detection — vsftpd Backdoor', level=2)
add_body(doc,
    'Running the Nmap vuln script against port 21 confirmed that vsftpd 2.3.4 contains the known '
    'backdoor vulnerability CVE-2011-2523. Nmap explicitly flagged it as VULNERABLE and exploitable.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_110509_flameshot.png',
    'Figure 11 — Nmap vuln script confirming CVE-2011-2523: vsftpd 2.3.4 backdoor VULNERABLE',
    width=Inches(5.5)
)

add_heading(doc, '3.5  FTP Anonymous Login Verification', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_110556_flameshot.png',
    'Figure 12 — FTP anonymous login confirmed on Metasploitable',
    width=Inches(5.5)
)

add_heading(doc, '3.6  SMTP User Enumeration (Port 25)', level=2)
add_body(doc,
    'smtp-user-enum was run against the SMTP service on port 25 using the VRFY method. '
    '20 valid system users were found.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_122415_flameshot.png',
    'Figure 13 — SMTP user enumeration: 20 valid users found on Metasploitable',
    width=Inches(5.5)
)

add_heading(doc, '3.7  RPC / NFS Enumeration (Port 111)', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/nmap/20260401_125641_flameshot.png',
    'Figure 14 — Nmap scan on port 111: rpcbind, NFS, mountd, nlockmgr exposed',
    width=Inches(5.5)
)

add_heading(doc, '3.8  Open Ports Summary', level=2)
add_table(doc,
    ['Port', 'Service', 'Version', 'Note'],
    [
        ['21',   'FTP',        'vsftpd 2.3.4',         'Backdoored — CVE-2011-2523'],
        ['22',   'SSH',        'OpenSSH 4.7p1',         'Outdated version'],
        ['23',   'Telnet',     'Linux telnetd',         'Cleartext protocol'],
        ['25',   'SMTP',       'Postfix smtpd',         '20 users enumerated'],
        ['53',   'DNS',        'ISC BIND 9.4.2',        'Exposed DNS service'],
        ['80',   'HTTP',       'Apache 2.2.8 (Ubuntu)', 'Version disclosed'],
        ['111',  'RPCBind',    '2 (RPC #100000)',        'NFS/mountd exposed'],
        ['139',  'SMB',        'Samba smbd 3.X–4.X',    'Exploitable version'],
        ['445',  'SMB',        'Samba smbd 3.X–4.X',    'Exploitable version'],
        ['1524', 'Bindshell',  'Root shell',             'Direct root access'],
        ['2049', 'NFS',        '2-4 (RPC #100003)',      'No access control'],
        ['3306', 'MySQL',      'MySQL 5.0.51a',          'No auth by default'],
        ['5432', 'PostgreSQL', 'PostgreSQL 8.3.0',       'Database exposed'],
        ['5900', 'VNC',        'Protocol 3.3',           'Remote desktop open'],
        ['6667', 'IRC',        'UnrealIRCd',             'Known vulnerable'],
        ['8180', 'HTTP',       'Apache Tomcat JSP 1.1',  'Management interface'],
    ],
    col_widths=[Inches(0.6), Inches(1.1), Inches(1.6), Inches(2.2)]
)

doc.add_page_break()

# ════════════════════════════════════════
# 4. TASK 2 — WIRESHARK
# ════════════════════════════════════════
add_heading(doc, '4. Task 2 — Network Traffic Analysis')
add_body(doc,
    'Wireshark was used to capture live network traffic on the vmnet1 (Host-Only) interface during '
    'active interaction with the target. The goal was to observe how different protocols behave at '
    'the packet level, identify cleartext credential exposure, and understand what information is '
    'visible to anyone on the same network segment.'
)

add_heading(doc, '4.1  HTTP Traffic Capture', level=2)
add_body(doc,
    'HTTP traffic between Kali and the Metasploitable web server (port 80) was captured. '
    'The packet inspection pane shows a POST request with form data — including field names and '
    'values — fully visible in plaintext. No decryption was needed.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_025229_flameshot.png',
    'Figure 15 — HTTP POST request captured: form data visible in cleartext',
    width=Inches(5.5)
)

add_heading(doc, '4.2  Traffic Overview — Multiple Protocols', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_030114_flameshot.png',
    'Figure 16 — Wireshark packet list: HTTP, TCP, and other protocol traffic visible',
    width=Inches(5.5)
)

add_heading(doc, '4.3  FTP Credentials in Plaintext', level=2)
add_body(doc,
    'FTP login packets were captured and inspected. The USER and PASS commands are sent in '
    'cleartext — the username "anonymous" and password "pass" are directly readable in '
    'the packet detail pane without any decryption.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_033412_flameshot.png',
    'Figure 17 — FTP USER anonymous captured in plaintext (Frame 188)',
    width=Inches(5.5)
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_033353_flameshot.png',
    'Figure 18 — FTP PASS captured in plaintext — password visible in packet (Frame 192)',
    width=Inches(5.5)
)

add_heading(doc, '4.4  Telnet Session — Unencrypted Terminal', level=2)
add_body(doc,
    'A Telnet connection from the target to the host was captured. The packet detail shows the '
    'Telnet negotiation (Terminal Type, Speed, X Display) in readable form. Every command in '
    'a Telnet session is transmitted as plaintext, making it trivial to read any session data.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_033327_flameshot.png',
    'Figure 19 — Telnet packet decoded: Terminal Type and Speed negotiation visible in Frame 251',
    width=Inches(5.5)
)

add_heading(doc, '4.5  RPC / Portmap Traffic', level=2)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/wireshark/20260402_034921_flameshot.png',
    'Figure 20 — RPC Portmap DUMP call captured (Frame 292) — NFS service enumeration',
    width=Inches(5.5)
)

add_heading(doc, '4.6  Filters Used', level=2)
add_table(doc,
    ['Filter Expression', 'Purpose', 'What Was Found'],
    [
        ['tcp.port == 21', 'FTP traffic',         'USER/PASS in plaintext'],
        ['http',           'HTTP requests',        'Form data, server headers'],
        ['tcp.port == 23', 'Telnet traffic',       'Full unencrypted session'],
        ['tcp.port == 111','RPC traffic',          'NFS/Portmap enumeration'],
        ['icmp',           'Ping verification',    'Host alive confirmation'],
    ],
    col_widths=[Inches(1.6), Inches(1.5), Inches(2.4)]
)

doc.add_page_break()

# ════════════════════════════════════════
# 5. TASK 3 — EXPLOITATION
# ════════════════════════════════════════
add_heading(doc, '5. Task 3 — Basic Exploitation')
add_body(doc,
    'The exploitation task targeted CVE-2011-2523 — a backdoor deliberately introduced into '
    'vsftpd 2.3.4 in 2011. When a username containing ":)" is sent to the FTP service, the '
    'backdoor opens a bind shell on TCP port 6200, giving the connecting client a root shell '
    'with no password required. The Metasploit Framework was used to automate this.'
)

add_heading(doc, '5.1  Vulnerability Details', level=2)
add_table(doc,
    ['Field', 'Detail'],
    [
        ['CVE ID',                  'CVE-2011-2523'],
        ['Affected Software',       'vsftpd 2.3.4'],
        ['Type',                    'Remote Code Execution — Backdoor'],
        ['Authentication Required', 'None'],
        ['Metasploit Module',       'exploit/unix/ftp/vsftpd_234_backdoor'],
        ['Result',                  'Root shell on target system'],
    ],
    col_widths=[Inches(2.2), Inches(3.3)]
)

add_heading(doc, '5.2  Exploitation Steps', level=2)
add_body(doc,
    'The full Metasploit session is shown below — from searching for the module through to '
    'receiving a root shell on the target.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/metasploitable/20260402_050826_flameshot.png',
    'Figure 21 — Metasploit: vsftpd_234_backdoor module configured and executed against 192.168.38.130',
    width=Inches(5.5)
)

add_heading(doc, '5.3  Root Shell Obtained', level=2)
add_body(doc,
    'After the exploit ran, a command shell session opened. The id command confirmed root access '
    '(uid=0). The ls command showed the full filesystem. The uname -a confirmed the target OS.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/loginshell/20260402_043208_flameshot.png',
    'Figure 22 — Root shell on Metasploitable: id = root, whoami = root, uname -a confirms OS',
    width=Inches(5.5)
)

add_heading(doc, '5.4  SSH Login to Metasploitable', level=2)
add_body(doc,
    'In addition to the exploit, SSH login was also tested using the default credentials '
    '(msfadmin / msfadmin) to confirm service access and further validate the attack surface.'
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/loginshell/20260402_051004_flameshot.png',
    'Figure 23 — SSH login to Metasploitable: Ubuntu 8.04 banner, successful authentication',
    width=Inches(5.5)
)
add_image(doc,
    '/home/fyxvoid/Void/ceh/assessment/screenshot/loginshell/20260402_051106_flameshot.png',
    'Figure 24 — SSH session active on Metasploitable as msfadmin',
    width=Inches(5.5)
)

doc.add_page_break()

# ════════════════════════════════════════
# 6. FINDINGS SUMMARY
# ════════════════════════════════════════
add_heading(doc, '6. Summary of Findings')
add_table(doc,
    ['#', 'Finding', 'Source', 'Severity'],
    [
        ['1', 'vsftpd 2.3.4 backdoor — unauthenticated root RCE (CVE-2011-2523)', 'Nmap + Metasploit', 'Critical'],
        ['2', 'Root bind shell open on port 1524',                                 'Nmap',             'Critical'],
        ['3', 'Samba 3.0.20 — known exploitable SMB version',                     'Nmap',             'Critical'],
        ['4', 'FTP transmits credentials in cleartext',                            'Wireshark',        'High'],
        ['5', 'Telnet active — full session visible without encryption',           'Wireshark + Nmap', 'High'],
        ['6', 'MySQL listening with no authentication required',                   'Nmap',             'High'],
        ['7', 'SMTP service exposed — 20 valid users enumerated',                  'smtp-user-enum',   'High'],
        ['8', 'Apache HTTP server version disclosed in response headers',          'Wireshark',        'Medium'],
        ['9', 'NFS share exposed with no access restrictions',                     'Nmap',             'Medium'],
        ['10','VNC remote desktop exposed on port 5900',                           'Nmap',             'Medium'],
    ],
    col_widths=[Inches(0.3), Inches(3.2), Inches(1.4), Inches(0.8)]
)

# ════════════════════════════════════════
# 7. CHALLENGES
# ════════════════════════════════════════
add_heading(doc, '7. Challenges Faced')

challenges = [
    ('Network Adapter Configuration',
     'VMware creates several virtual adapters and I initially selected the wrong one, which caused '
     'the ping to fail. Fixed by checking the VMware Virtual Network Editor and assigning the '
     'correct adapter to each VM.'),
    ('Metasploitable IP Identification',
     'Metasploitable uses DHCP by default. I had to log into the VM console and run ifconfig '
     'to get the assigned IP (192.168.38.130) before any scanning could start.'),
    ('Wireshark Interface Selection',
     'The host machine had multiple active interfaces. Selecting the wrong one captured no '
     'relevant traffic. Once the correct VMware vmnet1 adapter was selected, all lab traffic '
     'appeared as expected.'),
    ('Nmap Scan Speed',
     'Running -p- (all 65535 ports) was slow without rate control. Adding --min-rate 5000 '
     'brought the scan time down significantly while keeping the results accurate.'),
]

for title, text in challenges:
    p = doc.add_paragraph()
    r = p.add_run(title + ' — ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

# ════════════════════════════════════════
# 8. CONCLUSION
# ════════════════════════════════════════
add_heading(doc, '8. Conclusion')
add_body(doc,
    'This assessment covered the three core areas of practical security work. The lab was built '
    'on VMware Workstation using a Host-Only network, keeping all activity isolated. Nmap revealed '
    'a wide attack surface on the Metasploitable target — 16+ open services, several running '
    'dangerously outdated software with known CVEs.'
)
add_body(doc,
    'Wireshark confirmed how much sensitive information leaks over unencrypted protocols. FTP '
    'and Telnet both exposed credentials in plaintext — a critical finding in any real engagement. '
    'The exploitation phase demonstrated a full attack chain: the vsftpd 2.3.4 backdoor was '
    'exploited via Metasploit to obtain an unauthenticated root shell in under two minutes.'
)
add_body(doc,
    'All work was done within the isolated lab environment. This assessment provided a solid '
    'practical foundation in the tools and methodology used in real security assessments.'
)

# Declaration
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Declaration')
r.bold = True
r.font.size = Pt(12)

add_body(doc,
    'I confirm that this submission is my own work, completed as per WILA guidelines. '
    'All tasks were performed within the designated controlled lab environment.'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)

sig_table = doc.add_table(rows=2, cols=3)
sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ['Candidate Name', 'Signature', 'Date']
values = ['Sridharan T', '', 'April 2026']
for i, (lbl, val) in enumerate(zip(labels, values)):
    cell_lbl = sig_table.rows[0].cells[i]
    cell_val = sig_table.rows[1].cells[i]
    r = cell_lbl.paragraphs[0].add_run(lbl)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r2 = cell_val.paragraphs[0].add_run(val)
    r2.bold = True
    r2.font.size = Pt(11)
    # underline the value cell
    pPr = cell_val.paragraphs[0]._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot)
    pPr.append(pBdr)
    for c in [cell_lbl, cell_val]:
        c.width = Inches(2.0)

# ── Save ──
out = '/home/fyxvoid/Void/ceh/assessment/Sridharan_T_CommonAssessment.docx'
doc.save(out)
print(f'Saved: {out}')
